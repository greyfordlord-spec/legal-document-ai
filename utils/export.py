import os
from typing import Optional
from datetime import datetime
from docx import Document
from docx.shared import Inches
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch

class DocumentExporter:
    """Handles export of generated documents to PDF and DOCX formats."""
    
    def __init__(self, export_folder: str = "exports"):
        """Initialize exporter with export folder."""
        self.export_folder = export_folder
        self._ensure_export_folder()
    
    def _ensure_export_folder(self):
        """Ensure export folder exists."""
        if not os.path.exists(self.export_folder):
            os.makedirs(self.export_folder)
    
    def export_to_docx(self, content: str, document_type: str, language: str = "EN") -> Optional[str]:
        """
        Export document content to DOCX format.
        
        Args:
            content: Document content as string
            document_type: Type of document
            language: Language of document
            
        Returns:
            Path to exported file or None if failed
        """
        try:
            # Create document
            doc = Document()
            
            # Add title
            title = self._get_document_title(document_type, language)
            doc.add_heading(title, 0)
            
            # Add content
            paragraphs = content.split('\n\n')
            for paragraph in paragraphs:
                if paragraph.strip():
                    # Handle headers (lines starting with numbers)
                    if paragraph.strip().startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.')):
                        doc.add_heading(paragraph.strip(), level=1)
                    else:
                        doc.add_paragraph(paragraph.strip())
            
            # Generate filename
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"{document_type}_{language}_{timestamp}.docx"
            filepath = os.path.join(self.export_folder, filename)
            
            # Save document
            doc.save(filepath)
            return filepath
            
        except Exception as e:
            print(f"Error exporting to DOCX: {str(e)}")
            return None
    
    def export_to_pdf(self, content: str, document_type: str, language: str = "EN") -> Optional[str]:
        """
        Export document content to PDF format.
        
        Args:
            content: Document content as string
            document_type: Type of document
            language: Language of document
            
        Returns:
            Path to exported file or None if failed
        """
        try:
            # Generate filename
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"{document_type}_{language}_{timestamp}.pdf"
            filepath = os.path.join(self.export_folder, filename)
            
            # Create PDF document
            doc = SimpleDocTemplate(filepath, pagesize=A4 if language == "DE" else letter)
            styles = getSampleStyleSheet()
            
            # Create custom styles
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=16,
                spaceAfter=30,
                alignment=1  # Center alignment
            )
            
            normal_style = ParagraphStyle(
                'CustomNormal',
                parent=styles['Normal'],
                fontSize=11,
                spaceAfter=12
            )
            
            # Build story (content)
            story = []
            
            # Add title
            title = self._get_document_title(document_type, language)
            story.append(Paragraph(title, title_style))
            story.append(Spacer(1, 20))
            
            # Add content
            paragraphs = content.split('\n\n')
            for paragraph in paragraphs:
                if paragraph.strip():
                    # Handle headers (lines starting with numbers)
                    if paragraph.strip().startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.')):
                        story.append(Paragraph(paragraph.strip(), styles['Heading2']))
                    else:
                        story.append(Paragraph(paragraph.strip(), normal_style))
            
            # Build PDF
            doc.build(story)
            return filepath
            
        except Exception as e:
            print(f"Error exporting to PDF: {str(e)}")
            return None
    
    def _get_document_title(self, document_type: str, language: str) -> str:
        """Get document title based on type and language."""
        titles = {
            "residential_lease": {
                "EN": "Residential Lease Agreement",
                "DE": "Wohnungsmietvertrag"
            },
            "nda": {
                "EN": "Non-Disclosure Agreement",
                "DE": "Geheimhaltungsvereinbarung"
            },
            "b2b_contract": {
                "EN": "Business-to-Business Contract",
                "DE": "B2B-Vertrag"
            },
            "power_of_attorney": {
                "EN": "Power of Attorney",
                "DE": "Vollmacht"
            },
            "employment_contract": {
                "EN": "Employment Contract",
                "DE": "Arbeitsvertrag"
            },
            "meeting_minutes": {
                "EN": "Meeting Minutes",
                "DE": "Protokoll"
            }
        }
        
        return titles.get(document_type, {}).get(language, document_type)
    
    def get_export_formats(self) -> list[str]:
        """Get list of supported export formats."""
        return ["docx", "pdf"]
    
    def export_document(self, content: str, document_type: str, format_type: str, language: str = "EN") -> Optional[str]:
        """
        Export document in specified format.
        
        Args:
            content: Document content
            document_type: Type of document
            format_type: Export format (docx or pdf)
            language: Language of document
            
        Returns:
            Path to exported file or None if failed
        """
        if format_type.lower() == "docx":
            return self.export_to_docx(content, document_type, language)
        elif format_type.lower() == "pdf":
            return self.export_to_pdf(content, document_type, language)
        else:
            print(f"Unsupported export format: {format_type}")
            return None
